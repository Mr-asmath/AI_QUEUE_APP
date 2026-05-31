from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "doc" / "word" / "AI_Queue_App_Setup_Installation_Guide_Tharshini.docx"


def set_font(run, size=11, bold=False, color="000000", font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def heading(document, text, level=1):
    paragraph = document.add_heading(level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run, 16 if level == 1 else 13, True, "2E74B5" if level <= 2 else "1F4D78")


def para(document, text, bold_start=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    if bold_start and text.startswith(bold_start):
        run = paragraph.add_run(bold_start)
        set_font(run, bold=True, color="0B2545")
        run = paragraph.add_run(text[len(bold_start):])
        set_font(run)
    else:
        run = paragraph.add_run(text)
        set_font(run)


def bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_font(run)


def numbers(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_font(run)


def code(document, lines):
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(line)
        set_font(run, 9.5, False, "0B2545", "Courier New")


def table(document, rows, widths=None):
    doc_table = document.add_table(rows=1, cols=len(rows[0]))
    doc_table.style = "Table Grid"
    doc_table.autofit = False
    widths = widths or [Inches(1.8), Inches(4.7)]
    for index, cell_text in enumerate(rows[0]):
        cell = doc_table.rows[0].cells[index]
        if index < len(widths):
            cell.width = widths[index]
        run = cell.paragraphs[0].add_run(cell_text)
        set_font(run, bold=True, color="1F4D78")
    for row in rows[1:]:
        cells = doc_table.add_row().cells
        for index, cell_text in enumerate(row):
            if index < len(widths):
                cells[index].width = widths[index]
            run = cells[index].paragraphs[0].add_run(cell_text)
            set_font(run)
    document.add_paragraph()


def main():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.2

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("AI Queue App Setup and Installation Guide")
    set_font(run, 22, True, "0B2545")
    para(document, "Prepared for Tharshini to set up the project, work from GitHub, run Docker, and contribute safely.")

    heading(document, "Project Summary", 1)
    bullets(
        document,
        [
            "AI Queue App is a multi-industry queue management system.",
            "Frontend: React app served locally on port 3000.",
            "Backend: Flask API served locally on port 5000.",
            "Database/storage in Docker: SQLite app data volume, MongoDB for uploaded profile/logo assets, Redis cache, Prometheus metrics, mongo-express, and n8n automation.",
            "Main users: normal user, main admin, industry admin, queue operator, and service provider.",
        ],
    )

    heading(document, "Laptop Apps to Install", 1)
    table(
        document,
        [
            ["App", "Why it is needed"],
            ["Git", "Clone the GitHub repository, pull latest code, create branch, commit, and push changes."],
            ["GitHub account", "Access repository, push branch, and open pull requests."],
            ["Visual Studio Code", "Edit frontend, backend, Docker, Kubernetes, and documentation files."],
            ["Node.js 20 LTS", "Run React frontend locally and install npm packages."],
            ["Python 3.11", "Run Flask backend locally without Docker."],
            ["Docker Desktop", "Run frontend, backend, MongoDB, Redis, Prometheus, mongo-express, and n8n together."],
            ["Google Chrome or Microsoft Edge", "Test the running app in a browser."],
            ["MongoDB Compass optional", "Inspect MongoDB data with a desktop UI if needed."],
            ["Postman optional", "Test backend API routes."],
        ],
    )

    heading(document, "VS Code Extensions", 1)
    bullets(
        document,
        [
            "Python by Microsoft",
            "Pylance",
            "ESLint",
            "Prettier - Code formatter",
            "Docker",
            "GitLens optional",
            "YAML",
            "MongoDB for VS Code optional",
        ],
    )

    heading(document, "Project Packages", 1)
    para(document, "Frontend npm packages:", bold_start="Frontend")
    bullets(
        document,
        [
            "react 19.2.4",
            "react-dom 19.2.4",
            "react-scripts 5.0.1",
            "@testing-library/react, @testing-library/jest-dom, @testing-library/user-event, @testing-library/dom",
            "web-vitals",
        ],
    )
    para(document, "Backend Python packages:", bold_start="Backend")
    bullets(
        document,
        [
            "Flask 2.3.3",
            "Flask-CORS 4.0.0",
            "Flask-SQLAlchemy 3.0.5",
            "Werkzeug 2.3.8",
            "psycopg2-binary 2.9.6",
            "gunicorn 20.1.0",
            "python-dotenv 0.21.0",
            "redis 5.0.1",
            "pymongo 4.6.1",
        ],
    )
    para(document, "Docker services:", bold_start="Docker")
    bullets(
        document,
        [
            "backend: Flask API",
            "frontend: React build served by nginx",
            "mongodb: profile/logo asset storage",
            "redis: cache",
            "mongo-express: MongoDB browser UI",
            "prometheus: backend metrics",
            "n8n: automation workflows",
        ],
    )

    heading(document, "First Time Git Setup", 1)
    numbers(
        document,
        [
            "Install Git and sign in to GitHub.",
            "Ask the repository owner to add your GitHub username as collaborator.",
            "Open PowerShell or VS Code Terminal.",
            "Configure your Git name and email once.",
        ],
    )
    code(
        document,
        [
            'git config --global user.name "Tharshini"',
            'git config --global user.email "tharshini@example.com"',
            "git config --global init.defaultBranch main",
        ],
    )

    heading(document, "Clone the Repository", 1)
    para(document, "Replace the GitHub URL with the real repository URL that the owner shares.")
    code(
        document,
        [
            "cd E:\\Projects",
            "git clone https://github.com/<owner>/<repo-name>.git AI-Queue-App",
            "cd AI-Queue-App",
            "git status",
        ],
    )

    heading(document, "Branch Workflow for Tharshini", 1)
    para(document, "Always pull latest main first, then work in the tharshini branch.")
    code(
        document,
        [
            "git checkout main",
            "git pull origin main",
            "git checkout -b tharshini",
        ],
    )
    para(document, "If the branch already exists locally:")
    code(document, ["git checkout tharshini", "git pull origin tharshini"])
    para(document, "If the branch exists on GitHub but not locally:")
    code(document, ["git fetch origin", "git checkout -b tharshini origin/tharshini"])

    heading(document, "Daily Git Commands", 1)
    para(document, "Start of day: update main and rebase/merge into your branch.")
    code(
        document,
        [
            "git checkout main",
            "git pull origin main",
            "git checkout tharshini",
            "git merge main",
        ],
    )
    para(document, "After making changes: check, add, commit, and push.")
    code(
        document,
        [
            "git status",
            "git add .",
            'git commit -m "Update setup or feature work"',
            "git push origin tharshini",
        ],
    )
    para(document, "If this is the first push for the branch:")
    code(document, ["git push -u origin tharshini"])
    para(document, "To see changed files and history:")
    code(document, ["git status", "git diff", "git log --oneline --decorate -10"])

    heading(document, "Run With Docker", 1)
    para(document, "Recommended for a new laptop because Docker starts all services together.")
    numbers(
        document,
        [
            "Open Docker Desktop and wait until it says Docker is running.",
            "Open PowerShell or VS Code Terminal in the project root.",
            "Validate compose file.",
            "Build and run the stack.",
            "Open the app in the browser.",
        ],
    )
    code(
        document,
        [
            "cd E:\\Projects\\AI-Queue-App",
            "docker compose config",
            "docker compose up --build",
        ],
    )
    para(document, "Run in background:")
    code(document, ["docker compose up --build -d", "docker compose ps"])
    para(document, "Open these URLs:")
    bullets(
        document,
        [
            "Frontend: http://localhost:3000",
            "Backend health: http://localhost:5000/health",
            "Backend metrics: http://localhost:5000/metrics",
            "Mongo Express: http://localhost:8081",
            "Prometheus: http://localhost:9090",
            "n8n: http://localhost:5678",
        ],
    )
    para(document, "Useful Docker commands:")
    code(
        document,
        [
            "docker compose logs -f backend",
            "docker compose logs -f frontend",
            "docker compose restart backend",
            "docker compose restart frontend",
            "docker compose down",
            "docker compose down -v",
        ],
    )
    bullets(
        document,
        [
            "Use docker compose down to stop containers but keep data.",
            "Use docker compose down -v only when you want to delete local Docker volumes/data.",
        ],
    )

    heading(document, "Run Without Docker", 1)
    para(document, "Use this only if you want to debug frontend/backend directly.")
    para(document, "Backend terminal:")
    code(
        document,
        [
            "cd E:\\Projects\\AI-Queue-App\\backend",
            "python -m venv .venv",
            ".\\.venv\\Scripts\\activate",
            "python -m pip install --upgrade pip",
            "pip install -r requirements.txt",
            "python app.py",
        ],
    )
    para(document, "Frontend terminal:")
    code(
        document,
        [
            "cd E:\\Projects\\AI-Queue-App\\frontend",
            "npm install",
            "npm start",
        ],
    )

    heading(document, "Demo Accounts", 1)
    table(
        document,
        [
            ["Role", "Login"],
            ["Main admin", "admin@queue.com / admin123"],
            ["Industry admin", "industry@queue.com / demo123"],
            ["Queue operator", "operator@queue.com / demo123"],
            ["Provider", "provider@queue.com / demo123"],
        ],
    )

    heading(document, "Common Problems and Fixes", 1)
    bullets(
        document,
        [
            "If npm shows deprecated warnings, continue. Warnings are not errors.",
            "If npm fails with ETIMEDOUT, check internet, retry docker compose build frontend, or change network.",
            "If Docker says port already used, stop the other app using that port or edit docker-compose.yml.",
            "If login keeps failing after backend restart, clear browser cookies for localhost and try again.",
            "If images do not persist, confirm MongoDB container is running and volumes were not deleted.",
        ],
    )
    code(
        document,
        [
            "docker compose build frontend",
            "docker compose up -d",
            "docker compose logs -f frontend",
            "docker compose logs -f backend",
        ],
    )

    heading(document, "Before Pushing Work", 1)
    numbers(
        document,
        [
            "Run the app and manually check the page you changed.",
            "Check git status.",
            "Commit with a clear message.",
            "Push to tharshini branch.",
            "Tell the owner to review or open a pull request from tharshini into main.",
        ],
    )
    code(
        document,
        [
            "git status",
            "git add .",
            'git commit -m "Describe the change clearly"',
            "git push origin tharshini",
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
