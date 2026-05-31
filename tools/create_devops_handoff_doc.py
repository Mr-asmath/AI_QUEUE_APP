from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "doc" / "word" / "AI_Queue_App_DevOps_New_Joiner_Guide.docx"


def set_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run, size=16 if level == 1 else 13, bold=True, color="2E74B5" if level <= 2 else "1F4D78")
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_font(run)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_font(run)


def add_code_block(document, commands):
    for command in commands:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(command)
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string("0B2545")


def add_table(document, rows):
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.55), Inches(2.45), Inches(2.5)]
    for index, heading in enumerate(rows[0]):
        cell = table.rows[0].cells[index]
        cell.width = widths[index]
        run = cell.paragraphs[0].add_run(heading)
        set_font(run, bold=True, color="1F4D78")
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = widths[index]
            run = cells[index].paragraphs[0].add_run(value)
            set_font(run)
    document.add_paragraph()


def main():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("AI Queue App DevOps and New Joiner Guide")
    set_font(run, size=22, bold=True, color="0B2545")

    subtitle = document.add_paragraph()
    run = subtitle.add_run("Docker, Kubernetes, MongoDB assets, Redis cache, monitoring, and profile/session fixes")
    set_font(run, size=11, color="555555")

    add_heading(document, "What Changed", 1)
    add_bullets(
        document,
        [
            "Docker Compose now runs frontend, backend, MongoDB, Redis, Prometheus, mongo-express, and n8n with persistent volumes.",
            "Kubernetes manifests now include two replicas for the frontend and backend, health probes, PVC-backed app data, MongoDB, Redis, and Prometheus.",
            "Login sessions are persistent for 30 days by default, so users stay signed in until logout or cookie expiry.",
            "Redis caches the current-user response to reduce repeated login checks after page refreshes.",
            "Profile and logo image uploads are resized in the browser and stored in MongoDB when MONGODB_URI is configured.",
            "Browser geolocation permission prompts were removed, and the main-admin permission column was removed from the security screen.",
            "The backend exposes /metrics for Prometheus scraping and /health for container and Kubernetes probes.",
            "Registration terms now explain what user data is collected, why it is used, how sessions work, and what users/admins are responsible for.",
            "Frontend and backend loading or deployment downtime now shows a full updating page with retry timing and updating status.",
        ],
    )

    add_heading(document, "Service Map", 1)
    add_table(
        document,
        [
            ["Service", "Purpose", "Local port"],
            ["frontend", "React app served by nginx", "3000"],
            ["backend", "Flask API with gunicorn workers and threads", "5000"],
            ["mongodb", "Profile image and logo asset storage", "27017"],
            ["redis", "Current-user cache and fast repeated auth checks", "6379"],
            ["prometheus", "Scrapes backend metrics", "9090"],
            ["mongo-express", "Local MongoDB inspection UI", "8081"],
            ["n8n", "Automation workflow support", "5678"],
        ],
    )

    add_heading(document, "Run Locally With Docker", 1)
    add_numbered(
        document,
        [
            "From the repository root, run docker compose up --build.",
            "Open the frontend at http://localhost:3000 and the backend health endpoint at http://localhost:5000/health.",
            "Use http://localhost:9090 for Prometheus and http://localhost:8081 for mongo-express during local debugging.",
            "To reset local data, stop the stack and remove the named Docker volumes only when data loss is acceptable.",
        ],
    )

    add_heading(document, "Running Commands", 1)
    add_heading(document, "Docker Compose From Source", 2)
    add_code_block(
        document,
        [
            "cd E:\\Projects\\Live\\AI-Queue-App",
            "docker compose config",
            "docker compose up --build",
            "docker compose up --build -d",
            "docker compose ps",
            "docker compose logs -f backend",
            "docker compose logs -f frontend",
        ],
    )
    add_bullets(
        document,
        [
            "Use the non-detached command while debugging because logs stay visible.",
            "Use the -d command when running the stack in the background.",
        ],
    )

    add_heading(document, "Open Local Services", 2)
    add_code_block(
        document,
        [
            "Frontend:      http://localhost:3000",
            "Backend:       http://localhost:5000/health",
            "Metrics:       http://localhost:5000/metrics",
            "Prometheus:    http://localhost:9090",
            "Mongo Express: http://localhost:8081",
            "n8n:           http://localhost:5678",
        ],
    )

    add_heading(document, "Stop, Restart, and Rebuild", 2)
    add_code_block(
        document,
        [
            "docker compose stop",
            "docker compose start",
            "docker compose restart backend",
            "docker compose restart frontend",
            "docker compose down",
            "docker compose build --no-cache backend",
            "docker compose up -d backend",
        ],
    )

    add_heading(document, "Data Volumes", 2)
    add_code_block(
        document,
        [
            "docker volume ls",
            "docker compose down",
            "docker compose down -v",
        ],
    )
    add_bullets(
        document,
        [
            "docker compose down keeps named volumes and preserves app, MongoDB, Redis, Prometheus, and n8n data.",
            "docker compose down -v deletes named volumes. Use it only when a full local data reset is intended.",
        ],
    )

    add_heading(document, "Run Published GHCR Images", 2)
    add_code_block(
        document,
        [
            "cd E:\\Projects\\Live\\AI-Queue-App",
            "docker compose -f docker-compose.ghcr.yml config",
            "docker compose -f docker-compose.ghcr.yml pull",
            "docker compose -f docker-compose.ghcr.yml up -d",
            "docker compose -f docker-compose.ghcr.yml logs -f backend",
            "docker compose -f docker-compose.ghcr.yml down",
        ],
    )

    add_heading(document, "Kubernetes Deployment", 2)
    add_code_block(
        document,
        [
            "kubectl apply -f k8s/secrets.example.yaml",
            "kubectl apply -f k8s/mongodb.yaml",
            "kubectl apply -f k8s/redis.yaml",
            "kubectl apply -f k8s/backend.yaml",
            "kubectl apply -f k8s/frontend.yaml",
            "kubectl apply -f k8s/prometheus.yaml",
            "kubectl get pods",
            "kubectl get svc",
            "kubectl logs -f deploy/ai-queue-backend",
            "kubectl logs -f deploy/ai-queue-frontend",
            "kubectl rollout status deploy/ai-queue-backend",
            "kubectl rollout status deploy/ai-queue-frontend",
        ],
    )

    add_heading(document, "Kubernetes Local Port Forward", 2)
    add_code_block(
        document,
        [
            "kubectl port-forward svc/ai-queue-frontend 3000:80",
            "kubectl port-forward svc/ai-queue-backend 5000:5000",
            "kubectl port-forward svc/prometheus 9090:9090",
        ],
    )

    add_heading(document, "Kubernetes Cleanup", 2)
    add_code_block(
        document,
        [
            "kubectl delete -f k8s/prometheus.yaml",
            "kubectl delete -f k8s/frontend.yaml",
            "kubectl delete -f k8s/backend.yaml",
            "kubectl delete -f k8s/redis.yaml",
            "kubectl delete -f k8s/mongodb.yaml",
        ],
    )

    add_heading(document, "Local Verification Commands", 2)
    add_code_block(
        document,
        [
            "python -m py_compile backend\\app.py",
            "docker compose config",
            "docker compose -f docker-compose.ghcr.yml config",
            "cd frontend",
            "npm.cmd run build",
        ],
    )

    add_heading(document, "Deploy With Kubernetes", 1)
    add_numbered(
        document,
        [
            "Create secrets from k8s/secrets.example.yaml and replace default values before production use.",
            "Apply k8s/mongodb.yaml, k8s/redis.yaml, k8s/backend.yaml, k8s/frontend.yaml, and k8s/prometheus.yaml.",
            "Confirm pods are ready, then expose the frontend service or wire it into the cluster ingress.",
            "Keep SECRET_KEY stable across backend replicas so existing signed sessions remain valid.",
        ],
    )

    add_heading(document, "Important Notes", 1)
    add_bullets(
        document,
        [
            "The active backend is the Flask app in backend/app.py. It still uses SQLAlchemy for core queue records, while MongoDB is now used for uploaded profile/logo assets.",
            "For a full all-data MongoDB migration, plan a separate model and route migration from SQLAlchemy to MongoEngine or PyMongo collections.",
            "SQLite with two backend replicas is acceptable for local demos only. Production should use an external multi-writer database service or complete the MongoDB migration.",
            "Uploaded profile images are returned as data URLs from the API, so the avatar stays visible after refresh when the database and MongoDB volume are preserved.",
            "The app no longer calls navigator.geolocation and sends a Permissions-Policy header that disables geolocation.",
        ],
    )

    add_heading(document, "Terms and Conditions Added", 1)
    add_bullets(
        document,
        [
            "The /api/security/terms endpoint now returns real app terms for account, profile, queue, branch, security, audit, and uploaded image data.",
            "The terms explain that passwords are stored as secure hashes and sessions use protected cookies.",
            "The terms state that browser location, camera, microphone, contacts, files, and precise GPS are not collected by default.",
            "The terms describe how user data is used for queue operations, reports, notifications, support, monitoring, and security.",
            "The terms include user responsibilities, admin responsibilities, retention notes, service update availability, and limitation of use.",
        ],
    )

    add_heading(document, "Updating Page Behavior", 1)
    add_bullets(
        document,
        [
            "During initial frontend loading, the app shows a full-page Project is loading screen with the message FRONTEND LOADING.",
            "When the backend is unavailable, restarting, or updating, the app shows Project is updating with PLEASE TRY LATER.",
            "The backend updating screen displays an auto-retry countdown and a Try again now button.",
            "The bottom status text clearly shows Updating frontend... or Updating backend... so users know what is happening.",
            "The app checks /health every 30 seconds while the backend is unavailable and returns to normal once the backend is healthy.",
        ],
    )

    add_heading(document, "New Joiner Checklist", 1)
    add_bullets(
        document,
        [
            "Read README.md and RUNNING.md for existing project behavior.",
            "Run python -m py_compile backend/app.py after backend edits.",
            "Run docker compose config after infrastructure edits.",
            "Run npm.cmd run build from frontend on Windows, because PowerShell can block npm.ps1.",
            "Check /health and /metrics before opening a pull request or handing off a deployment.",
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
