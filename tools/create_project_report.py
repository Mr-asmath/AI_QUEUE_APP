from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "doc" / "word"
OUT_PATH = OUT_DIR / "AI_Queue_App_Project_Report.docx"


MANUAL_NOTES = r'''modify the project in make an ai agent automation like use n8n or any free agent now. the token can generate the users avoid the duplications. modify the application is used for multiple place like hospital, schools, banks, hotals and more. so there have main admin that mean i, manage and control all on application. then the users or industry admin can install or open the application to send the request to access the application and fill the details of admin, industry and more for get the access from main admin. the main admin can reject to send the message or email this is reject. or the main admin is accept to send the admin username and default generate password to industry admin. and send the term and condition of all the details can view and manage from main admin(application admin). the industry admin can login or reset the password. then the industry admin can create a branches and allocate the queue managers for each branches. to create a branch ask a branch name, details, what type of branch like (hospital, school, hotal, bank and more and add the lst of other(requied the type name)) on first site, then second site to ask the what are include in queue dashboard the users site, queue operator and service provider. the users can generate the token anywhere, to give the details, to see the previews suggestion details, current queue count. the queue operator can see the current queue count, see the suggestion and details on user or previews users. the service provider can role like( doctor on hospital, teacher or admin on school, manager on hostal, manager on bank) like it. the industry admin can first fill the service provider site dashboard to display the more fields in checkbox type like (display user details, display perviews users details, display next details, display current user details, display suggestion text box (to ask the number of suggestion text box and ask title for each suggestion box), display search checkbox (to allow to add the items for display on checkbox the service provider can select multiple item, then ask checkbox value like price, name, or any value for checkbox which is only needed to users), display the cash prize(to diplay the checkbox items with prize to select the items on checkbox lastly display the total, average, count and give more aggerations, the aggearation also selected from creator), display the transations(send, recevied, pending, failure like display for users this is used for online or offline transation on industry entris),to display the ai suggestion textbox( it is generate the suggestions based on user queue current details and generate the suggestion to display on suggestion textbox, this is work after the service provider click ask suggestion button) and add like this more based on multiple industries. second to allowd to create the queue operator site, in the side to add to display above items and ask each on two options edit and read( allow edit to able to edit the details, read mean only read the details), to allow the queue operator for allocate the service provider to user based on users need and service provider role, thrid allow to create an user site to create a key and values( key means names of attributes, value means value of key there is select the type like int, text, float, image, or more types, then add the checkbox for display the items like ( diplay previews suggestions, display current suggestions, display current queue count, display to allow generate the queue token, display to allow the reject queue token, display cash prize, display transations, display queue operator name and contact, display service provider name and contacts, display total number of queue count and like more. the industry admin can allow to create a multiplt queue operator and service provider. the user can download or display on website to allow on register and login, the users can select the place, select the industry, fill the requird details, then click submit, then queue operator can verified the details and generate the token to send the message on user to see on application. to avoid the duplicate users. each token can valide only 30 to 60 min only within time the users can entry on industry and view the queue operator, the queue operator can verified the click customer in,the token can valid, otherwise user not entry before 60 min the token is automatically reject and send the notification to user eacn 10 min before verifications. the users can queue order if any users not on industry to displaying pending to reorder the pending user in after next custom in automatically. the industry admin can create a branches to display select the allow option in two options like to ask allow all branch, allow to selected branch only for share the details of users and all for user can use same industry in multiple places with display the same details and suggestions for all dashboard sites. to add the ai agent for control the all automation works, and automatically detect the duplicate tokens, to generate the ai sugestions based on user inputs on current token. and create with more classical style like real time application to use normal database to allow for access the cloud db in future options, to add the docker for manage the software and add the kubernets with 2 replicasets. and add more devops tools to manage this project. to modify the project for the main admin can create and add the industry admin not allow to login , the indusrty admin can register the page the admin can allow only create a account otherwise see waiting, or admin can reject to cancel the registation. the operates and provider so register in separate register with role type this is allow to create an account from that currect industry admin. there is display the unique number can each industry generate by industry admin to create a branch , the operator or provider is select that unique code to send the request to that current industry admin this is all work on industry admin. and don't display the admin login email adn password on login page and add the register page for other 3 actor


==========

to modify the project for industry admin dashboard site to allow to change the password, to add the forgot password option to click forgot password to send on admin the admin can click accept to send the default password to industry admin adn same to 3 roles. and operator and servicer is create there are is not generate to default password. the branch is create option display this Create Branch
Branch name
Details

Hospital
Suggestion
Consultation:300, Service Charge:100
name:text, phone:text, need:text  i need to create like follow this to create a branch ask a branch name, details, what type of branch like (hospital, school, hotal, bank and more and add the lst of other(requied the type name)) on first site, then second site to ask the what are include in queue dashboard the users site, queue operator and service provider. the users can generate the token anywhere, to give the details, to see the previews suggestion details, current queue count. the queue operator can see the current queue count, see the suggestion and details on user or previews users. the service provider can role like( doctor on hospital, teacher or admin on school, manager on hostal, manager on bank) like it. the industry admin can first fill the service provider site dashboard to display the more fields in checkbox type like (display user details, display perviews users details, display next details, display current user details, display suggestion text box (to ask the number of suggestion text box and ask title for each suggestion box), display search checkbox (to allow to add the items for display on checkbox the service provider can select multiple item, then ask checkbox value like price, name, or any value for checkbox which is only needed to users), display the cash prize(to diplay the checkbox items with prize to select the items on checkbox lastly display the total, average, count and give more aggerations, the aggearation also selected from creator), display the transations(send, recevied, pending, failure like display for users this is used for online or offline transation on industry entris),to display the ai suggestion textbox( it is generate the suggestions based on user queue current details and generate the suggestion to display on suggestion textbox, this is work after the service provider click ask suggestion button) and add like this more based on multiple industries. second to allowd to create the queue operator site, in the side to add to display above items and ask each on two options edit and read( allow edit to able to edit the details, read mean only read the details), to allow the queue operator for allocate the service provider to user based on users need and service provider role, thrid allow to create an user site to create a key and values( key means names of attributes, value means value of key there is select the type like int, text, float, image, or more types, then add the checkbox for display the items like ( diplay previews suggestions, display current suggestions, display current queue count, display to allow generate the queue token, display to allow the reject queue token, display cash prize, display transations, display queue operator name and contact, display service provider name and contacts, display total number of queue count and like more. the industry admin can allow to create a multiplt queue operator and service provider. the user can download or display on website to allow on register and login, the users can select the place, select the industry, fill the requird details, then click submit, then queue operator can verified the details and generate the token to send the message on user to see on application. to avoid the duplicate users. each token can valide only 30 to 60 min only within time the users can entry on industry and view the queue operator, the queue operator can verified the click customer in,the token can valid, otherwise user not entry before 60 min the token is automatically reject and send the notification to user eacn 10 min before verifications. the users can queue order if any users not on industry to displaying pending to reorder the pending user in after next custom in automatically. the industry admin can create a branches to display select the allow option in two options like to ask allow all branch, allow to selected branch only for share the details of users and all for user can use same industry in multiple places with display the same details and suggestions for all dashboard sites. you can follow this all. the create branch and all side you add the checkbox method not typing method




can you see the whole project for any modifications, add styles, animations, images. to make a real app. with access on any devices, the header side displaying like box in profile and access to change like top header bar side profile is display like circle default face to add the user and other for change the more default logos, face or own images.



in industryadmin page, to modify the top side apps and your profile remove and app is convert to home logo on leftside and your profile change to click on project logo or image, remove the logout button on dashboard add on inside of project page in separate the sections. in dashboard not display the create branch section to change on plus logo in right botton to click to move to create branch page to create. in branch section ony display the created branchs same for staff and queue control. and correct the algin and change checkbox to display like listbox to click one or more to add on below box for only checkbosx inputs. to modify and add the address  and more personal details for create the queue manager and queue controller. in Your profile page i ask project avater logo not like letters, the images get like input type image not like url image. like upload image for all profiles

in industry admin page the home button is change to home logo and move to right side of title lable side, the plue logo is click to display is ok but there is not have terminator button or cancel button to add the both and terminator button like x in right top side. and in aggeration field not understanding give clearly there is displaying only v a o correctly and project page add the side navigation for like menu for account details, app logo, change password, account session. first display the account details to click to display that only section. the change password is ok but add forgot password also this is send to main add to create a default password and send to industry add email.

in main admin page to add the user management page for view all types of users in separatly and they company, branchs, works, details and more. to add the user id in 6 digit for all it is used to mainly app users to use the id to find on any branch in same company ti fetch the details, to add the search bar for search name,branch, number, id

modify the queue operator page for add the display of history of queues and details to add the search bar to find the user to display the user details and more. add the date from and to options also. and separate the currect queue and queue history in separate menu options, the queue operator in queue history displaying in vertical i need like table field format roms and column and there is displaying two search bar remove this Find user
Search user name, phone, branch, email, or 6-digit ID. in samething changge all records in table format not like box columns . in service provider page ai can give the suggest first so modify the ai to add the text box the service provided can text the requirements to click the ai suggestion to generate not generate the first if not click that to add doctor suggest requirements only or click ask suggestion to send both doctor and ai suggestions'''


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    set_table_borders(t)
    if widths is None:
        widths = [9360 // len(headers)] * len(headers)
    for idx, header in enumerate(headers):
        cell = t.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            r.font.size = Pt(8.5)
    doc.add_paragraph()
    return t


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_code_block(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.style = "Code"
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(31, 45, 61)


def add_callout(doc, title, body):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    set_table_borders(t, "CBD5E1")
    cell = t.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run("\n" + body)
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(85, 85, 85)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        s = styles[style_name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)

    if "Code" not in styles:
        code = styles.add_style("Code", 1)
        code.font.name = "Consolas"
        code.font.size = Pt(8.5)
        code.paragraph_format.left_indent = Inches(0.15)
        code.paragraph_format.space_after = Pt(0)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "AI Queue Automation Project Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("AI_QUEUE_APP | Project Documentation")
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def add_title_page(doc):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("AI Queue Automation Application")
    p2 = doc.add_paragraph(style="Subtitle")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Full Project Details, Use Cases, Operations, DevOps, and Manual Notes")
    doc.add_paragraph()
    meta_rows = [
        ("Repository", "https://github.com/Mr-asmath/AI_QUEUE_APP"),
        ("Prepared for", "Mr. Asmath"),
        ("Prepared on", date.today().strftime("%B %d, %Y")),
        ("Local workspace", str(ROOT)),
        ("Primary application type", "Multi-industry AI-assisted queue automation platform"),
    ]
    table(doc, ["Field", "Value"], meta_rows, [2200, 7160])
    add_callout(
        doc,
        "Document Purpose",
        "This report records the current project architecture, user flows, run commands, deployment configuration, GitHub Actions, tools, services, use cases, and the complete source manual notes used to guide project development."
    )
    doc.add_page_break()


def add_abstract(doc):
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "AI Queue Automation is a multi-tenant queue management application designed for hospitals, schools, banks, hotels, and other service industries. It supports centralized main-admin control, industry-admin onboarding, configurable branches, queue managers, service providers, app users, token generation, duplicate-token prevention, queue history, AI-assisted provider suggestions, notifications, Dockerized execution, n8n automation, Kubernetes manifests, GitHub Actions CI, GitHub Container Registry images, and GitHub Pages frontend deployment."
    )
    doc.add_paragraph(
        "The system is implemented as a React frontend and Flask backend with SQLAlchemy models and SQLite by default, while keeping a path open for cloud database migration. The application is structured around role-based workflows: main admin, industry admin, queue operator/manager, service provider/controller, and user."
    )
    doc.add_paragraph("Keywords: AI queue automation, multi-industry platform, token management, n8n, Docker, Kubernetes, GitHub Actions, React, Flask, SQLAlchemy.")


def add_project_overview(doc):
    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "The project addresses queue control for service organizations where users need to request service, receive a valid token, be verified by a queue operator, and be routed to the appropriate service provider. The system is intended to reduce duplicate requests, provide configurable dashboards for different industries, preserve historical queue records, and support automation for reminders and token expiry."
    )
    table(
        doc,
        ["Layer", "Main Files", "Purpose"],
        [
            ("Frontend", "frontend/src/App.js, frontend/src/components/*, frontend/src/App.css", "React application, dashboards, authentication screens, profile management, queue views, and responsive UI."),
            ("Backend", "backend/app.py", "Flask API, role authorization, models, token lifecycle, automation endpoint, profile/password workflows, queue history."),
            ("Automation", "automation/n8n-workflow.md, /api/automation/run", "n8n schedule trigger calls backend automation to expire old tokens and send reminders."),
            ("AI Agent", "ai-agent/*.py and backend make_ai_suggestion", "Priority, wait prediction, queue optimization helpers, and provider-triggered suggestion generation."),
            ("DevOps", "Dockerfile, docker-compose.yml, k8s/*.yaml, .github/workflows/*.yml", "Container builds, local orchestration, Kubernetes replicas, CI, GHCR publish, and Pages deployment."),
        ],
        [1700, 3300, 4360],
    )
    add_callout(
        doc,
        "Current Public Frontend Target",
        "The GitHub Pages frontend target is https://mr-asmath.github.io/AI_QUEUE_APP/. A live backend must be hosted separately and configured through the REACT_APP_API_ORIGIN GitHub repository variable."
    )


def add_roles_use_cases(doc):
    doc.add_heading("2. Actors and Use Cases", level=1)
    table(
        doc,
        ["Actor", "Primary Responsibilities", "Key Use Cases"],
        [
            ("Main Admin", "Application owner and global controller.", "Approve/reject industry access, view user directory, handle default-password requests, monitor industries/branches/users."),
            ("Industry Admin", "Controls one industry/company.", "Create branches, configure dashboards, create staff, manage project profile/logo/password/session, view queues and history."),
            ("Queue Operator / Queue Manager", "Controls queue verification and routing.", "View current queue, verify users, mark customer-in, allocate providers, reject tokens, search history by date/name/branch/id."),
            ("Service Provider / Queue Controller", "Completes service work.", "View allocated users, enter provider requirements, ask AI suggestion, select checkbox items/prices, complete service."),
            ("User / App User", "Requests service from any allowed branch.", "Register/login, select industry/place/branch, fill dynamic details, generate token, view queue count, notifications, suggestions."),
            ("n8n Automation Agent", "Runs scheduled backend automation.", "POST automation endpoint every 5 minutes to expire tokens and send reminders."),
        ],
        [1600, 3300, 4460],
    )
    doc.add_heading("2.1 Representative Use Cases", level=2)
    for item in [
        "A hospital patient requests a token, enters need/phone/name details, and receives queue status.",
        "A school admin branch uses the same queue engine for student/parent service workflows.",
        "A bank queue manager verifies the user's identity and routes the token to the appropriate counter manager.",
        "A hotel manager uses service provider checkbox items and amount aggregations for service charges.",
        "A main admin approves a new company request and sends generated credentials to the industry admin.",
        "A service provider types requirements, clicks Ask AI Suggestion, and completes service with a combined provider/AI note.",
    ]:
        add_bullet(doc, item)


def add_architecture(doc):
    doc.add_heading("3. Architecture", level=1)
    doc.add_paragraph("The architecture uses a browser-based React application, a Flask REST API, a relational database layer, optional n8n automation, and containerized deployment. The design is intentionally simple for local development but includes paths for container registry, Kubernetes, and cloud database migration.")
    add_code_block(
        doc,
        """
User Browser
  -> React Frontend (localhost:3000, nginx container, or GitHub Pages)
  -> Flask Backend API (localhost:5000 or hosted backend URL)
  -> SQLAlchemy database (SQLite now, PostgreSQL-ready through DATABASE_URL)
  -> Notifications, token lifecycle, queue history, suggestions
  -> n8n Scheduler / Automation Agent
        -> POST /api/automation/run every 5 minutes
        -> expire old tokens and send reminders
"""
    )
    table(
        doc,
        ["Component", "Technology", "Notes"],
        [
            ("Frontend", "React 19, react-scripts, CSS", "Single-page dashboard app with role-based rendering and responsive pages."),
            ("Backend API", "Flask, Flask-CORS, Flask-SQLAlchemy, Werkzeug", "Session-based auth, role checks, token APIs, admin APIs, automation."),
            ("Database", "SQLite default; psycopg2-binary included", "DATABASE_URL supports later cloud PostgreSQL migration."),
            ("Runtime", "Gunicorn in backend Docker image", "Container command binds Flask app to 0.0.0.0:5000."),
            ("Automation", "n8n self-hosted container", "Free automation workflow calls protected backend endpoint."),
            ("Orchestration", "Docker Compose and Kubernetes", "Compose for local stack; Kubernetes manifests use 2 replicas."),
            ("CI/CD", "GitHub Actions, GHCR, GitHub Pages", "Build checks, Docker publish, frontend static deployment."),
        ],
        [1800, 2400, 5160],
    )


def add_data_api(doc):
    doc.add_heading("4. Data Model and API Surface", level=1)
    table(
        doc,
        ["Entity", "Important Fields", "Function"],
        [
            ("User", "id, user_code, name, email, role, phone, address, designation, emergency_contact, personal_details, avatar fields", "Stores all actors, including 6-digit user IDs for cross-branch lookup."),
            ("Industry", "name, type, admin_id, logo fields", "Company/industry tenant controlled by an industry admin."),
            ("Branch", "name, details, branch_type, dashboard_config, user_schema", "Configurable service location with dynamic fields and dashboard options."),
            ("Token", "token_code, token_number, status, user/industry/branch/provider/operator ids, details_json, duplicate_key, expires_at", "Queue request lifecycle record."),
            ("Suggestion", "token_id, provider_id, suggestion_text, selected_items_json, aggregates_json", "Completed service notes, selected checkbox items, totals/average/count."),
            ("Notification", "user_id, message, type, token_id, is_read", "In-app messages for users and admins."),
            ("PasswordResetRequest", "user_id, admin_id, status, generated_password", "Admin-approved default password reset flow."),
        ],
        [1600, 4300, 3460],
    )
    doc.add_heading("4.1 API Catalogue", level=2)
    table(
        doc,
        ["Area", "Endpoints"],
        [
            ("Auth", "/api/auth/register, /api/auth/login, /api/auth/logout, /api/auth/me, /api/auth/forgot-password, /api/auth/request-default-password, /api/auth/reset-password/<token>"),
            ("Profile", "/api/profile"),
            ("Main Admin", "/api/admin/access-requests, /api/admin/users/directory, /api/admin/password-reset-requests"),
            ("Industry Admin", "/api/industry/branches, /api/industry/staff, /api/industry/users/search"),
            ("Queue", "/api/token, /api/queue/status, /api/operator/queue, /api/operator/queue-history, /api/operator/tokens/<id>/action"),
            ("Provider", "/api/provider/tokens, /api/provider/tokens/<id>/ai-suggestion, /api/provider/tokens/<id>/suggestion"),
            ("User", "/api/user/my-tokens, /api/user/my-suggestions, /api/user/notifications"),
            ("Automation", "/api/automation/run"),
            ("Health", "/health"),
        ],
        [1700, 7660],
    )


def add_user_workflows(doc):
    doc.add_heading("5. Step-by-Step User Workflows", level=1)
    workflows = {
        "Main Admin": [
            "Login as application admin.",
            "Open access requests.",
            "Review industry/admin details submitted by a requester.",
            "Approve to create an industry admin with a generated default password, or reject with a message.",
            "Open user management to review all roles, companies, branches, work records, and personal details.",
            "Review password reset requests and approve default-password generation when needed.",
        ],
        "Industry Admin": [
            "Login with generated credentials or reset password.",
            "Open profile/project page to manage account details, logo/avatar, password, forgot-password request, and session/logout.",
            "Open branches list to view created branches.",
            "Click the plus/create action to open branch creation.",
            "Enter branch name, details, type, optional other type name, user schema, and dashboard options.",
            "Configure service provider, queue operator, and user dashboard options using checkbox/listbox style controls.",
            "Create queue managers and queue controllers with address, designation, emergency contact, and personal details.",
        ],
        "Queue Operator / Manager": [
            "Open Current Queue menu.",
            "Review active tokens, user details, branch, status, and suggestions.",
            "Verify token details.",
            "Click Customer In when the user arrives within valid time.",
            "Allocate the token to a service provider based on need and role.",
            "Open Queue History menu to search rows by name, branch, phone, email, 6-digit ID, token, status, date from, and date to.",
        ],
        "Service Provider / Controller": [
            "Open allocated service tokens.",
            "Review user details and branch details.",
            "Type provider requirements or service notes.",
            "Optionally click Ask AI Suggestion; AI is not generated first.",
            "Select one or more configured checkbox items such as consultation or service charge.",
            "Review total, average, and count aggregations.",
            "Complete service with provider-only notes or provider plus AI suggestion.",
        ],
        "App User": [
            "Register or login.",
            "Select industry/place and branch.",
            "Fill required dynamic branch fields.",
            "Submit token request.",
            "Avoid duplicate tokens by using active token status instead of submitting again.",
            "View current queue count, token status, notifications, and previous/current suggestions.",
        ],
    }
    for role, steps in workflows.items():
        doc.add_heading(role, level=2)
        for step in steps:
            add_number(doc, step)


def add_ai_automation(doc):
    doc.add_heading("6. AI and Automation", level=1)
    doc.add_paragraph("The project contains two automation concepts: a backend-integrated suggestion generator and a free self-hosted n8n workflow that invokes the backend automation endpoint.")
    table(
        doc,
        ["Feature", "Implementation", "Behavior"],
        [
            ("Duplicate token guard", "Token duplicate_key plus active status checks", "Prevents repeated active token creation for the same user/branch/details while the token is still valid."),
            ("Token validity", "TOKEN_TTL_MINUTES, default 60", "Tokens expire when users do not arrive within the configured validity period."),
            ("Reminder window", "REMINDER_WINDOW_MINUTES, default 10", "Automation can notify users before expiry."),
            ("AI suggestion", "make_ai_suggestion(token, provider_requirements)", "Provider types requirements first; AI suggestion is generated only on button click."),
            ("AI service helpers", "ai-agent/main.py, priority.py, prediction.py, optimizer.py", "Priority scoring, wait prediction, completion prediction, and queue optimization concepts."),
            ("n8n workflow", "automation/n8n-workflow.md", "Schedule Trigger every 5 minutes calls POST /api/automation/run."),
        ],
        [1800, 3300, 4260],
    )
    doc.add_heading("6.1 n8n Automation Steps", level=2)
    for step in [
        "Open http://localhost:5678.",
        "Login using admin / admin123.",
        "Create a workflow with Schedule Trigger every 5 minutes.",
        "Add an HTTP Request node.",
        "Set Method to POST.",
        "Set URL to http://backend:5000/api/automation/run.",
        "Set header X-Automation-Key: local-automation-key.",
        "Activate the workflow.",
    ]:
        add_number(doc, step)


def add_running_commands(doc):
    doc.add_heading("7. Running Commands", level=1)
    doc.add_heading("7.1 Demo Accounts", level=2)
    add_code_block(
        doc,
        """
Main admin:      admin@queue.com / admin123
Industry admin:  industry@queue.com / demo123
Queue operator:  operator@queue.com / demo123
Provider:        provider@queue.com / demo123
"""
    )
    doc.add_heading("7.2 Run Locally Without Docker", level=2)
    add_code_block(
        doc,
        """
# Terminal 1 - Backend
cd backend
python -m pip install --upgrade pip
pip install -r requirements.txt
python app.py

# Terminal 2 - Frontend
cd frontend
npm install
npm start
"""
    )
    doc.add_heading("7.3 Health Checks", level=2)
    add_code_block(
        doc,
        """
Invoke-WebRequest http://localhost:5000/health -UseBasicParsing
Invoke-WebRequest http://localhost:5000/api/catalog -UseBasicParsing
"""
    )
    doc.add_heading("7.4 Docker Compose", level=2)
    add_code_block(
        doc,
        """
docker compose up --build

# Services
Frontend: http://localhost:3000
Backend:  http://localhost:5000
n8n:      http://localhost:5678

# Stop
docker compose down

# Stop and delete volumes/database
docker compose down -v

# Logs
docker compose logs -f backend
docker compose logs -f frontend
docker compose logs -f n8n
"""
    )
    doc.add_heading("7.5 Docker Images from GitHub Container Registry", level=2)
    add_code_block(
        doc,
        """
docker compose -f docker-compose.ghcr.yml config
docker compose -f docker-compose.ghcr.yml up

# From frontend folder, helper file also exists:
cd frontend
docker compose -f docker-compose.ghcr.yml config
"""
    )
    doc.add_heading("7.6 Docker Build Per Service", level=2)
    add_code_block(
        doc,
        """
docker build -t ai-queue-backend ./backend
docker run --name ai-queue-backend --rm -p 5000:5000 `
  -e SECRET_KEY=change-me `
  -e DATABASE_URL=sqlite:////data/multi_industry_queue.db `
  -e TOKEN_TTL_MINUTES=60 `
  -e REMINDER_WINDOW_MINUTES=10 `
  -v ai-queue-data:/data `
  ai-queue-backend

docker build -t ai-queue-frontend ./frontend
docker run --name ai-queue-frontend --rm -p 3000:80 ai-queue-frontend
"""
    )
    doc.add_heading("7.7 Kubernetes", level=2)
    add_code_block(
        doc,
        """
docker build -t ai-queue-backend:latest ./backend
docker build -t ai-queue-frontend:latest ./frontend

kubectl apply -f k8s/secrets.example.yaml
kubectl apply -f k8s/backend.yaml
kubectl apply -f k8s/frontend.yaml

kubectl get pods
kubectl get services

kubectl port-forward service/ai-queue-backend 5000:5000
kubectl port-forward service/ai-queue-frontend 3000:80

kubectl delete -f k8s/frontend.yaml
kubectl delete -f k8s/backend.yaml
kubectl delete -f k8s/secrets.example.yaml
"""
    )


def add_git_devops(doc):
    doc.add_heading("8. Git, GitHub, GitHub Actions, and Deployment", level=1)
    doc.add_paragraph("The Git remote is configured for the repository https://github.com/Mr-asmath/AI_QUEUE_APP.git. The main branch is used for workflow triggers.")
    table(
        doc,
        ["Workflow", "File", "Purpose"],
        [
            ("CI and Docker Publish", ".github/workflows/ci-deploy.yml", "Checks backend compile/import, builds frontend, publishes backend and frontend Docker images to GHCR on push."),
            ("Deploy Frontend to GitHub Pages", ".github/workflows/github-pages.yml", "Builds React frontend with PUBLIC_URL=/AI_QUEUE_APP and deploys frontend static files to GitHub Pages."),
        ],
        [2200, 3300, 3860],
    )
    doc.add_heading("8.1 Git Commands", level=2)
    add_code_block(
        doc,
        """
git status
git add .
git commit -m "Project update"
git push origin main
git log --oneline -5
git remote -v
"""
    )
    doc.add_heading("8.2 GitHub Pages", level=2)
    add_code_block(
        doc,
        """
Frontend URL:
https://mr-asmath.github.io/AI_QUEUE_APP/

GitHub settings:
Settings -> Pages -> Source -> GitHub Actions

Backend URL variable:
Repository -> Settings -> Secrets and variables -> Actions -> Variables
REACT_APP_API_ORIGIN=https://your-backend-url
"""
    )
    doc.add_heading("8.3 GitHub Container Registry", level=2)
    add_code_block(
        doc,
        """
Backend image:
ghcr.io/mr-asmath/ai-queue-app-backend:latest

Frontend image:
ghcr.io/mr-asmath/ai-queue-app-frontend:latest
"""
    )


def add_tools_services(doc):
    doc.add_heading("9. Tools and Services", level=1)
    table(
        doc,
        ["Tool / Service", "Role in Project"],
        [
            ("React", "Frontend dashboard and user interface."),
            ("react-scripts", "Development server and production build pipeline."),
            ("Flask", "Backend API service."),
            ("Flask-CORS", "Cross-origin frontend/backend access."),
            ("Flask-SQLAlchemy", "ORM for users, industries, branches, tokens, suggestions, notifications."),
            ("SQLite", "Default local database."),
            ("PostgreSQL-ready dependency", "psycopg2-binary supports future cloud database use via DATABASE_URL."),
            ("Gunicorn", "Production Python server inside backend Docker image."),
            ("Docker", "Build and run backend/frontend containers."),
            ("Docker Compose", "Run backend, frontend, n8n, and volumes locally."),
            ("n8n", "Free self-hosted automation scheduler."),
            ("Kubernetes", "Deployment manifests with two replicas for backend and frontend."),
            ("GitHub Actions", "Build, verify, publish, and deploy workflows."),
            ("GitHub Container Registry", "Stores published Docker images."),
            ("GitHub Pages", "Hosts static React frontend over HTTPS."),
            ("PowerShell", "Windows command environment for local operations."),
        ],
        [2600, 6760],
    )


def add_testing_troubleshooting(doc):
    doc.add_heading("10. Testing, Validation, and Troubleshooting", level=1)
    doc.add_heading("10.1 Validation Commands", level=2)
    add_code_block(
        doc,
        """
python -m py_compile backend/app.py
python -c "import sys; sys.path.insert(0, 'backend'); import app; print('backend import ok')"

cd frontend
npm run build

docker compose -f docker-compose.ghcr.yml config
"""
    )
    doc.add_heading("10.2 Common Problems", level=2)
    table(
        doc,
        ["Problem", "Cause", "Fix"],
        [
            ("Docker Desktop Linux engine missing", "Docker Desktop is not running or Linux containers are off.", "Open Docker Desktop, wait for running state, run docker version and docker info."),
            ("Docker build path not found", "Command run from nested backend/frontend folder with wrong context.", "Use docker build -t image . inside service folder, or ./backend from project root."),
            ("Port already in use", "Existing Python, Node, or Docker process using 3000/5000.", "Inspect process list and stop the conflicting process."),
            ("GitHub Pages Not Found error", "Pages not enabled or configure-pages call fails.", "Set Settings -> Pages -> Source -> GitHub Actions; workflow now deploys without configure-pages."),
            ("Frontend loads but login fails on Pages", "Backend URL is still localhost.", "Set REACT_APP_API_ORIGIN to the hosted backend URL and rerun Pages workflow."),
        ],
        [2200, 3300, 3860],
    )


def add_requirements_traceability(doc):
    doc.add_heading("11. Requirements Traceability", level=1)
    table(
        doc,
        ["Requirement Area", "Implemented / Planned Response"],
        [
            ("Multi-industry use", "Branch and industry types cover hospital, school, hotel, bank, and other with custom type name."),
            ("Main admin control", "Access request approval/rejection, user directory, password reset request management."),
            ("Industry admin branch creation", "Create branches with dashboard config, schema, branch type, details, plus/cancel/terminator style UI."),
            ("Staff creation", "Queue manager/provider creation with personal and address details."),
            ("Duplicate tokens", "Duplicate key and active-token guard prevent repeated valid tokens."),
            ("Queue history", "Separate Current Queue and Queue History menu, table rows/columns, search and date filters."),
            ("AI suggestion flow", "Provider types requirements first; AI suggestion generated only after click."),
            ("Profile image/logo", "Profile page supports avatar/logo presets and upload-style image input behavior in project UI."),
            ("DevOps", "Docker, Docker Compose, GitHub Actions, GHCR, GitHub Pages, Kubernetes manifests."),
            ("Automation", "n8n workflow calls backend automation endpoint for token expiry/reminders."),
        ],
        [2600, 6760],
    )


def add_appendix_manual_notes(doc):
    doc.add_page_break()
    doc.add_heading("Appendix A. Manual Notes Preserved Verbatim", level=1)
    add_callout(
        doc,
        "Preservation Note",
        "The following section preserves the user's manual notes as source requirements. Spelling, wording, capitalization, and line breaks are intentionally retained."
    )
    for block in MANUAL_NOTES.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        if block.strip() == "==========":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(block)
            r.bold = True
        elif not block:
            p.add_run(" ")
        else:
            r = p.add_run(block)
            r.font.size = Pt(9)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)
    add_title_page(doc)
    add_abstract(doc)
    add_project_overview(doc)
    add_roles_use_cases(doc)
    add_architecture(doc)
    add_data_api(doc)
    add_user_workflows(doc)
    add_ai_automation(doc)
    add_running_commands(doc)
    add_git_devops(doc)
    add_tools_services(doc)
    add_testing_troubleshooting(doc)
    add_requirements_traceability(doc)
    add_appendix_manual_notes(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
